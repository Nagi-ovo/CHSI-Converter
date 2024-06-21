from flask import Flask, request, render_template, send_from_directory, make_response
from werkzeug.utils import secure_filename
import os
from add_float_picture import add_float_picture
from extract_img import extract_image_from_pdf
from extract_info import extract_info_from_pdf
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
import uuid
import shutil  # 确保引入shutil模块

app = Flask(__name__)

def convert_to_docx(path):
    try:
        extracted_info = extract_info_from_pdf(path)
        doc = Document("static/template.docx")

        paragraph = doc.add_paragraph()
        doc.element.body.insert(1, paragraph._element)
        paragraph.alignment = 1
        paragraph.add_run('Update date:' + extracted_info['Update Date'])

        del extracted_info['Update Date']

        table = doc.add_table(rows=1, cols=2)
        table.autofit = False

        for cell in table.columns[0].cells:
            cell.width = Inches(0.5)
        for cell in table.columns[1].cells:
            cell.width = Inches(5.0)

        border_xml = '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">' \
                    '<w:top w:val="nil"/>' \
                    '<w:left w:val="nil"/>' \
                    '<w:bottom w:val="nil"/>' \
                    '<w:right w:val="nil"/>' \
                    '</w:tcBorders>'

        for key, value in extracted_info.items():
            cells = table.add_row().cells
            for cell in cells:
                cell._element.get_or_add_tcPr().append(parse_xml(border_xml))
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            is_last = key == list(extracted_info.keys())[-1]
            cells[0].text = key + ("" if is_last else "\n") 
            cells[1].text = value + ("" if is_last else "\n")

        cropped_image_1 = extract_image_from_pdf(path, 1, 1898, 583, 2230, 1026)
        add_float_picture(doc.add_paragraph(), cropped_image_1, width=Inches(1.2), pos_x=Pt(430), pos_y=Pt(140))

        cropped_image_2 = extract_image_from_pdf(path, 1, 300, 2690, 630, 2985)
        add_float_picture(doc.add_paragraph(), cropped_image_2, width=Inches(1.2), pos_x=Pt(78), pos_y=Pt(643))

        output_path = path.replace(".pdf", ".docx")
        doc.save(output_path)

        return output_path
    
    except Exception as e:
        return make_response(f"<script>alert('Error during DOCX conversion: {e}'); window.location.href = document.referrer;</script>")

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert_file():
    if 'file' not in request.files:
        return make_response("<script>alert('缺少文件部分'); window.location.href = document.referrer;</script>")

    file = request.files['file']
    if file.filename == '':
        return make_response("<script>alert('没有选中的文件'); window.location.href = document.referrer;</script>")

    if not file.filename.lower().endswith('.pdf'):
        return make_response("<script>alert('只接受 PDF 文件'); window.location.href = document.referrer;</script>")

    if not file.filename.startswith('教育部学籍在线验证报告_'):
        return make_response("<script>alert('请不要传入无关文件'); window.location.href = document.referrer;</script>")
    
    try:
        filename = secure_filename(file.filename)
        filepath = os.path.join(os.getcwd(), 'upload', filename)
        file.save(filepath)

        output_path = convert_to_docx(filepath)

        directory = os.path.dirname(output_path)
        filename = os.path.basename(output_path)
        output_filename = str(uuid.uuid4()) + '.docx'

        response = make_response(send_from_directory(directory, filename, as_attachment=True))
        response.headers["Content-Disposition"] = f"attachment; filename={output_filename}"
        
        # 隐私处理
        upload_folder = os.path.join(os.getcwd(), 'upload')
        for filename in os.listdir(upload_folder):
            if filename != '.gitkeep': 
                file_path = os.path.join(upload_folder, filename)
                try:
                    if os.path.isfile(file_path) or os.path.islink(file_path):
                        os.unlink(file_path)
                    elif os.path.isdir(file_path):
                        shutil.rmtree(file_path)
                except Exception as e:
                    print(f"Failed to delete {file_path}. Reason: {e}")
        
        return response
    except Exception as e:
        return make_response(f"<script>alert('处理文件时发生错误: {e}'); window.location.href = document.referrer;</script>")

if __name__ == '__main__':
    app.run(debug=True, port=5001, host='0.0.0.0')
else:
    application=app
