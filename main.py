from tkinter import Tk, Button, Label, StringVar
from tkinter.filedialog import askopenfilename
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml import OxmlElement  # Add this line
import os
import subprocess
import sys
from extract_info import extract_info_from_pdf
from extract_img import extract_image_from_pdf
from add_float_picture import add_float_picture



# TODO:unable to move the table left
def set_table_position(table, left):
    tblpPr = OxmlElement('w:tblpPr')
    tblpPr_left = OxmlElement('w:leftFromText')
    tblpPr_left.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w'] = str(left)
    tblpPr.append(tblpPr_left)
    table._element.tblPr.append(tblpPr)


def open_docx_on_desktop(file_path):
    if sys.platform.startswith('darwin'):  # macOS
        command = ['open', file_path]
    elif sys.platform.startswith('win32'):  # Windows
        os.startfile(file_path)
    elif sys.platform.startswith('linux'):  # Linux
        command = ['xdg-open', file_path]
    else:
        print("无法识别的操作系统")
        return

    if sys.platform.startswith('linux') or sys.platform.startswith('darwin'):
        try:
            subprocess.run(command, check=True)
        except subprocess.CalledProcessError as e:
            print("无法打开文件:", e)


root = Tk()
root.title("EZ4STU")
root.geometry("600x400")
root.configure(bg='lightblue')

title_label = Label(root, text="Xuexin Converter", font=("Helvetica", 24, "bold"), bg='lightblue', fg='white')
title_label.pack(pady=20)

file_path = StringVar()
file_path.set("请选择要转换的PDF文件")

converted_file_path = StringVar()
converted_file_path.set("")

def select_file():
    path = askopenfilename(filetypes=[("PDF Files", "*.pdf")])
    file_path.set(path)

select_button = Button(root, text="选择文件", command=select_file, font=("Arial", 16), bg='white', fg='black')
select_button.pack(pady=10)

file_label = Label(root, textvariable=file_path, font=("Arial", 14), bg='lightblue', fg='white')
file_label.pack(pady=10)


def convert_to_docx():
    path = file_path.get()
    if path:
        extracted_info = extract_info_from_pdf(path)
        doc = Document("template.docx")

        

        table = doc.add_table(rows=1, cols=2)
        table.autofit = False

        set_table_position(table, 5440)  # Move table 0.5 inch from the left margin

        for cell in table.columns[0].cells:
            cell.width = Inches(0.5)
        for cell in table.columns[1].cells:
            cell.width = Inches(5.0)

        border_xml = '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">' \
                    '<w:top w:val="nil"/>' \
                    '<w:left w:val="nil"/>' \
                    '<w:bottom w:val="nil"/>' \
                    '<w:right w:val="nil"/>' \
                    '</w:tcBorders>'

        for key, value in extracted_info.items():
            cells = table.add_row().cells
            for cell in cells:
                cell._element.get_or_add_tcPr().append(parse_xml(border_xml))
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            cells[0].text = key
            cells[1].text = value + "\n"

        #add_float_picture(doc.add_paragraph(), cropped_image_2, width=Inches(1.2), pos_x=Pt(430), pos_y=Pt(140))

        cropped_image_1 = extract_image_from_pdf(path, 1, 1898, 583, 2230, 1026)
        add_float_picture(doc.add_paragraph(), cropped_image_1, width=Inches(1.2), pos_x=Pt(430), pos_y=Pt(140))

        cropped_image_2 = extract_image_from_pdf(path, 1, 300, 2690, 630, 2985)
        add_float_picture(doc.add_paragraph(), cropped_image_2, width=Inches(1.2), pos_x=Pt(78), pos_y=Pt(642))
        
        output_path = path.replace(".pdf", ".docx")
        doc.save(output_path)

        converted_file_path.set("已转换为DOCX文件：\n" + output_path)

        root.after(1500, lambda: open_docx_on_desktop(output_path))

    else:
        file_path.set("请先选择一个PDF文件")

convert_button = Button(root, text="转换", command=convert_to_docx, font=("Arial", 16), bg='white', fg='black')
convert_button.pack(pady=20)

converted_file_label = Label(root, textvariable=converted_file_path, font=("Arial", 14), bg='lightblue', fg='white')
converted_file_label.pack(pady=10)

root.mainloop()
