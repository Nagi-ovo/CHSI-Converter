from flask import request, make_response, send_from_directory
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, register_element_cls
from docx.oxml.ns import nsdecls
from docx.oxml.shape import CT_Picture
from docx.oxml.xmlchemy import BaseOxmlElement, OneAndOnlyOne
from pypdf import PdfReader
from pypinyin import lazy_pinyin
from pdf2image import convert_from_path
import re
import os
import uuid
import shutil


class CT_Anchor(BaseOxmlElement):
    extent = OneAndOnlyOne("wp:extent")
    docPr = OneAndOnlyOne("wp:docPr")
    graphic = OneAndOnlyOne("a:graphic")

    @classmethod
    def new(cls, cx, cy, shape_id, pic, pos_x, pos_y):
        anchor = parse_xml(cls._anchor_xml(pos_x, pos_y))
        anchor.extent.cx = cx
        anchor.extent.cy = cy
        anchor.docPr.id = shape_id
        anchor.docPr.name = f"Picture {shape_id}"
        anchor.graphic.graphicData.uri = (
            "http://schemas.openxmlformats.org/drawingml/2006/picture"
        )
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def new_pic_anchor(cls, shape_id, rId, filename, cx, cy, pos_x, pos_y):
        pic_id = 0
        pic = CT_Picture.new(pic_id, filename, rId, cx, cy)
        anchor = cls.new(cx, cy, shape_id, pic, pos_x, pos_y)
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def _anchor_xml(cls, pos_x, pos_y):
        return (
            '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" \n'
            '           behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1" \n'
            "           %s>\n"
            '  <wp:simplePos x="0" y="0"/>\n'
            '  <wp:positionH relativeFrom="page">\n'
            "    <wp:posOffset>%d</wp:posOffset>\n"
            "  </wp:positionH>\n"
            '  <wp:positionV relativeFrom="page">\n'
            "    <wp:posOffset>%d</wp:posOffset>\n"
            "  </wp:positionV>\n"
            '  <wp:extent cx="914400" cy="914400"/>\n'
            "  <wp:wrapNone/>\n"
            '  <wp:docPr id="666" name="unnamed"/>\n'
            "  <wp:cNvGraphicFramePr>\n"
            '    <a:graphicFrameLocks noChangeAspect="1"/>\n'
            "  </wp:cNvGraphicFramePr>\n"
            "  <a:graphic>\n"
            '    <a:graphicData uri="URI not set"/>\n'
            "  </a:graphic>\n"
            "</wp:anchor>" % (nsdecls("wp", "a", "pic", "r"), int(pos_x), int(pos_y))
        )


def new_pic_anchor(part, image_descriptor, width, height, pos_x, pos_y):
    rId, image = part.get_or_add_image(image_descriptor)
    cx, cy = image.scaled_dimensions(width, height)
    shape_id, filename = part.next_id, image.filename
    return CT_Anchor.new_pic_anchor(shape_id, rId, filename, cx, cy, pos_x, pos_y)


class CHSIConverter:
    @staticmethod
    def extract_text_from_pdf(pdf_path):
        try:
            with open(pdf_path, "rb") as pdf_file_obj:
                pdf_reader = PdfReader(pdf_file_obj)
                text = "".join(page.extract_text() for page in pdf_reader.pages)
            return text
        except Exception as e:
            return make_response(
                f"<script>alert('从 PDF 提取文本错误：{e}'); window.location.href = document.referrer;</script>"
            )

    @staticmethod
    def extract_info(patterns_dict, text):
        results = {}
        for prop, pattern in patterns_dict.items():
            match = re.search(pattern, text)
            if match:
                value = match.group(1)
                if prop == "Name":
                    pinyin = lazy_pinyin(value)
                    results[prop] = (
                        f"{''.join(pinyin[1:]).capitalize()} {pinyin[0].capitalize()}"
                    )
                elif prop == "Gender":
                    results[prop] = (
                        "Male" if value == "男" else "Female" if value == "女" else None
                    )
                elif prop == "Ethnic":
                    results[prop] = "".join(lazy_pinyin(value[:-1])).title()
                elif prop in [
                    "Date of Birth",
                    "Date of Enrollment",
                    "Update Date",
                    "Expected Graduation Date",
                ]:
                    year, rest = value.split("年")
                    month, day = rest.split("月")
                    results[prop] = f"{month}/{day.replace('日', '')}/{year}"
                elif prop == "Levels":
                    results[prop] = "Undergraduate"
                elif prop == "Form":
                    results[prop] = "General full-time remote study"
                elif prop == "Educational System":
                    results[prop] = f"{value} years"
                elif prop == "Type":
                    results[prop] = "General higher education"
                elif prop == "School Status":
                    results[prop] = "Currently enrolled"
                else:
                    results[prop] = value
            else:
                results[prop] = None
        return results

    @classmethod
    def extract_info_from_pdf(cls, path):
        text = cls.extract_text_from_pdf(path)
        patterns_dict = {
            prop: re.compile(r"{}\s*([^\s]*)".format(pattern))
            for prop, pattern in {
                "Update Date": "更新日期：",
                "Name": "姓名",
                "Gender": "性别",
                "Id Number": "证件号码",
                "Ethnic": "民族",
                "Date of Birth": "出生日期 ",
                "Institution": "院校",
                "Levels": "层次",
                "Faculties": "院系",
                "Class": "班级",
                "Major": "专业",
                "Student Number": "学号",
                "Form": "学习形式",
                "Date of Enrollment": "入学日期",
                "Educational System": "学制",
                "Type": "学历类别",
                "School Status": "学籍状态",
                "Expected Graduation Date": "预计毕业日期",
            }.items()
        }
        extracted_info = cls.extract_info(patterns_dict, text)

        return extracted_info

    @staticmethod
    def extract_image_from_pdf(path, page_number, left, top, right, bottom):
        try:
            images = convert_from_path(
                path, dpi=300, first_page=page_number, last_page=page_number
            )
            image = images[0]
            cropped_image = image.crop((left, top, right, bottom))
            file_name = os.path.splitext(path)[0]
            image_path = f"{file_name}_image.png"
            cropped_image.save(image_path)
            return image_path
        except Exception as e:
            return make_response(
                f"<script>alert('从 PDF 提取图片错误：{e}'); window.location.href = document.referrer;</script>"
            )

    @staticmethod
    def add_float_picture(
        p, image_path_or_stream, width=None, height=None, pos_x=0, pos_y=0
    ):
        try:
            run = p.add_run()
            anchor = new_pic_anchor(
                run.part, image_path_or_stream, width, height, pos_x, pos_y
            )
            run._r.add_drawing(anchor)
        except Exception:
            return make_response(
                "<script>alert('浮动图片添加时发生错误'); window.location.href = document.referrer;</script>"
            )

    @classmethod
    def convert_to_docx(cls, path):
        try:
            extracted_info = cls.extract_info_from_pdf(path)
            if not isinstance(extracted_info, dict):
                return make_response(
                    f"<script>alert('Error extracting PDF info.'); window.location.href = document.referrer;</script>"
                )

            doc = Document("static/template.docx")

            # Add update date paragraph
            paragraph = doc.add_paragraph()
            doc.element.body.insert(1, paragraph._element)
            paragraph.alignment = 1
            run = paragraph.add_run("Update date:" + extracted_info["Update Date"])
            run.font.name = "Times New Roman"

            extracted_info.pop("Update Date", None)

            # Add table immediately after update date
            table = doc.add_table(rows=0, cols=2)
            doc.element.body.insert(2, table._element)
            table.autofit = False

            for cell in table.columns[0].cells:
                cell.width = Inches(0.3)
            for cell in table.columns[1].cells:
                cell.width = Inches(5.2)

            border_xml = (
                '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:top w:val="nil"/>'
                '<w:left w:val="nil"/>'
                '<w:bottom w:val="nil"/>'
                '<w:right w:val="nil"/>'
                "</w:tcBorders>"
            )

            for key, value in extracted_info.items():
                cells = table.add_row().cells
                for cell in cells:
                    cell._element.get_or_add_tcPr().append(parse_xml(border_xml))
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph = cell.paragraphs[0]
                    paragraph.style.font.name = "Times New Roman"

                is_last = key == list(extracted_info.keys())[-1]
                cells[0].text = key + ("" if is_last else "\n")
                cells[1].text = str(value) + ("" if is_last else "\n")

            cropped_image_1 = cls.extract_image_from_pdf(path, 1, 1898, 583, 2230, 1026)
            cls.add_float_picture(
                doc.add_paragraph(),
                cropped_image_1,
                width=Inches(1.2),
                pos_x=Pt(430),
                pos_y=Pt(140),
            )

            cropped_image_2 = cls.extract_image_from_pdf(path, 1, 300, 2690, 630, 2985)
            cls.add_float_picture(
                doc.add_paragraph(),
                cropped_image_2,
                width=Inches(1.2),
                pos_x=Pt(78),
                pos_y=Pt(643),
            )

            output_path = path.replace(".pdf", ".docx")
            doc.save(output_path)

            return output_path

        except Exception as e:
            return make_response(
                f"<script>alert('Error during DOCX conversion: {e}'); window.location.href = document.referrer;</script>"
            )

    @classmethod
    def convert_file(cls):
        if "file" not in request.files:
            return make_response(
                "<script>alert('缺少文件部分'); window.location.href = document.referrer;</script>"
            )

        file = request.files["file"]
        if file.filename == "":
            return make_response(
                "<script>alert('没有选中的文件'); window.location.href = document.referrer;</script>"
            )

        if not file.filename.lower().endswith(".pdf"):
            return make_response(
                "<script>alert('只接受 PDF 文件'); window.location.href = document.referrer;</script>"
            )

        try:
            filename = secure_filename(file.filename)
            filepath = os.path.join(os.getcwd(), "upload", filename)
            file.save(filepath)

            output_path_or_response = cls.convert_to_docx(filepath)

            if not output_path_or_response or not isinstance(
                output_path_or_response, str
            ):
                if (
                    isinstance(output_path_or_response, tuple)
                    and len(output_path_or_response) > 1
                ):
                    return output_path_or_response
                return make_response(
                    f"<script>alert('DOCX conversion failed.'); window.location.href = document.referrer;</script>"
                )

            output_path = output_path_or_response

            directory = os.path.dirname(output_path)
            filename = os.path.basename(output_path)
            output_filename = str(uuid.uuid4()) + ".docx"

            response = make_response(
                send_from_directory(directory, filename, as_attachment=True)
            )
            response.headers["Content-Disposition"] = (
                f"attachment; filename={output_filename}"
            )

            upload_folder = os.path.join(os.getcwd(), "upload")
            for filename in os.listdir(upload_folder):
                if filename != ".gitkeep":
                    file_path = os.path.join(upload_folder, filename)
                    try:
                        if os.path.isfile(file_path) or os.path.islink(file_path):
                            os.unlink(file_path)
                        elif os.path.isdir(file_path):
                            shutil.rmtree(file_path)
                    except Exception as e:
                        print(f"Failed to delete {file_path}. Reason: {e}")

            return response
        except Exception as e:
            return make_response(
                f"<script>alert('处理文件时发生错误：{e}'); window.location.href = document.referrer;</script>"
            )


register_element_cls("wp:anchor", CT_Anchor)
